"""
Document Parser Module
Handles reading and extracting text from DOC/DOCX files for job descriptions and resumes.
"""

import os
from docx import Document
from typing import Optional
import PyPDF2
import pdfplumber
import io


class DocumentParser:
    """A utility class to parse DOC/DOCX and PDF files and extract text content."""
    
    @staticmethod
    def read_docx_file(file_path: str) -> Optional[str]:
        """
        Read and extract text from a DOCX file.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            Optional[str]: Extracted text content or None if error
        """
        try:
            if not os.path.exists(file_path):
                print(f"Error: File not found - {file_path}")
                return None
                
            if not file_path.lower().endswith(('.docx', '.doc')):
                print(f"Error: Unsupported file format - {file_path}")
                return None
                
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            return '\n'.join(text_content)
            
        except Exception as e:
            print(f"Error reading file {file_path}: {str(e)}")
            return None
    
    @staticmethod
    def read_pdf_file(file_path: str) -> Optional[str]:
        """
        Read and extract text from a PDF file.
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            Optional[str]: Extracted text content or None if error
        """
        try:
            if not os.path.exists(file_path):
                print(f"Error: File not found - {file_path}")
                return None
                
            text_content = []
            
            # Try pdfplumber first (better for complex layouts)
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content.append(page_text.strip())
                
                if text_content:
                    return '\n'.join(text_content)
            
            except Exception:
                # Fallback to PyPDF2
                text_content = []
                
            # Fallback method using PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_content.append(page_text.strip())
            
            return '\n'.join(text_content) if text_content else None
            
        except Exception as e:
            print(f"Error reading PDF file {file_path}: {str(e)}")
            return None
    
    @staticmethod
    def read_file(file_path: str) -> Optional[str]:
        """
        Read and extract text from a file (supports DOCX, DOC, and PDF).
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            Optional[str]: Extracted text content or None if error
        """
        if not os.path.exists(file_path):
            print(f"Error: File not found - {file_path}")
            return None
        
        file_extension = file_path.lower()
        
        if file_extension.endswith(('.docx', '.doc')):
            return DocumentParser.read_docx_file(file_path)
        elif file_extension.endswith('.pdf'):
            return DocumentParser.read_pdf_file(file_path)
        else:
            print(f"Error: Unsupported file format - {file_path}")
            return None
    
    @staticmethod
    def validate_file(file_path: str) -> bool:
        """
        Validate if the file exists and has supported format.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            bool: True if file is valid, False otherwise
        """
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return False
            
        if not file_path.lower().endswith(('.docx', '.doc', '.pdf')):
            print(f"Unsupported file format: {file_path}. Supported formats: .docx, .doc, .pdf")
            return False
            
        return True
    
    @staticmethod
    def read_job_description(file_path: str) -> Optional[str]:
        """
        Read job description from a DOC/DOCX/PDF file.
        
        Args:
            file_path (str): Path to the job description file
            
        Returns:
            Optional[str]: Job description text or None if error
        """
        print(f"Reading job description from: {file_path}")
        return DocumentParser.read_file(file_path)
    
    @staticmethod
    def read_resume(file_path: str) -> Optional[str]:
        """
        Read resume content from a DOC/DOCX/PDF file.
        
        Args:
            file_path (str): Path to the resume file
            
        Returns:
            Optional[str]: Resume text or None if error
        """
        print(f"Reading resume from: {file_path}")
        return DocumentParser.read_file(file_path)